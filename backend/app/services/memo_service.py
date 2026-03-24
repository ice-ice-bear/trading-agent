"""투자 메모 생성 서비스"""
import json
import logging
from datetime import datetime
from app.models.db import execute_query, execute_insert

logger = logging.getLogger(__name__)


async def generate_memo_html(signal_id: int) -> str | None:
    """시그널 데이터 기반 HTML 투자 메모 생성"""
    rows = await execute_query("SELECT * FROM signals WHERE id = ?", (signal_id,))
    if not rows:
        return None

    signal = dict(rows[0])
    scenarios = json.loads(signal.get("scenarios_json") or "{}")
    dart = json.loads(signal.get("dart_fundamentals_json") or "{}")
    expert_stances = json.loads(signal.get("expert_stances_json") or "{}")
    metadata = json.loads(signal.get("metadata_json") or "{}")

    direction_kr = {"buy": "매수", "sell": "매도", "hold": "보유"}.get(signal.get("direction", ""), signal.get("direction", ""))
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    rr_score = signal.get("rr_score", 0) or 0

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>{signal.get('stock_name', '')} ({signal.get('stock_code', '')}) 투자 메모</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Noto Sans KR', sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #333; line-height: 1.6; }}
  h1 {{ border-bottom: 2px solid #333; padding-bottom: 8px; }}
  .meta {{ color: #666; font-size: 0.9em; }}
  .section {{ margin: 24px 0; }}
  .section h2 {{ color: #1a1a1a; font-size: 1.2em; border-left: 4px solid #007bff; padding-left: 12px; }}
  table {{ border-collapse: collapse; width: 100%; margin-top: 8px; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  th {{ background: #f5f5f5; }}
  .positive {{ color: #22c55e; }}
  .negative {{ color: #ef4444; }}
</style>
</head>
<body>
<h1>{signal.get('stock_name', '')} ({signal.get('stock_code', '')})</h1>
<p class="meta">생성일: {generated_at} | 시그널 #{signal_id} | R/R Score: {rr_score:.1f}</p>

<div class="section">
<h2>투자 의견</h2>
<p><strong>{direction_kr}</strong></p>
"""

    if signal.get("variant_view"):
        html += f'<p><strong>시장 오해:</strong> {signal["variant_view"]}</p>'

    html += '</div>'

    # 시나리오
    if scenarios:
        html += '<div class="section"><h2>시나리오 분석</h2><table><tr><th>시나리오</th><th>목표가</th><th>상승률</th><th>확률</th></tr>'
        for key in ["bull", "base", "bear"]:
            s = scenarios.get(key, {})
            if s:
                upside = s.get("upside_pct", 0) or 0
                prob = s.get("probability", 0) or 0
                price = s.get("price_target", 0) or 0
                css_class = "positive" if upside >= 0 else "negative"
                html += f'<tr><td>{s.get("label", key)}</td><td>{price:,.0f}원</td><td class="{css_class}">{upside:+.1f}%</td><td>{prob * 100:.0f}%</td></tr>'
        html += '</table></div>'

    # DART 펀더멘탈
    if dart:
        html += '<div class="section"><h2>펀더멘탈 (DART)</h2><table><tr><th>지표</th><th>값</th></tr>'
        labels = {
            "dart_per": "PER", "dart_pbr": "PBR", "dart_operating_margin": "영업이익률 (%)",
            "dart_debt_ratio": "부채비율 (%)", "dart_eps_yoy_pct": "EPS 성장률 (%)",
            "dart_dividend_yield": "배당수익률 (%)",
        }
        for key, label in labels.items():
            val = dart.get(key)
            if val is not None:
                formatted = f"{val:.1f}" if isinstance(val, (int, float)) else str(val)
                html += f'<tr><td>{label}</td><td>{formatted}</td></tr>'
        html += '</table></div>'

    # 전문가 패널
    if expert_stances:
        html += '<div class="section"><h2>전문가 패널</h2><table><tr><th>전문가</th><th>의견</th></tr>'
        for name, stance in expert_stances.items():
            html += f'<tr><td>{name}</td><td>{stance}</td></tr>'
        html += '</table></div>'

    # 메타데이터 (수급, 내부자, DCF 등)
    dcf = metadata.get("dcf_valuation")
    if dcf and dcf.get("fair_value"):
        html += f'<div class="section"><h2>DCF 적정가</h2><p>적정가: <strong>{dcf["fair_value"]:,.0f}원</strong>'
        if dcf.get("upside_pct") is not None:
            html += f' ({"+" if dcf["upside_pct"] >= 0 else ""}{dcf["upside_pct"]:.1f}%)'
        html += '</p></div>'

    investor = metadata.get("investor_trend")
    if investor:
        html += f'<div class="section"><h2>수급 동향 ({investor.get("days", 20)}일)</h2>'
        html += f'<p>외국인 순매수: {investor.get("foreign_net_buy", 0):+,}주 | 기관 순매수: {investor.get("institution_net_buy", 0):+,}주</p></div>'

    html += '<hr><p class="meta">본 메모는 AI 분석 시스템에 의해 자동 생성되었습니다. 투자 판단의 참고 자료로만 활용하세요.</p>'
    html += '</body></html>'

    # DB 기록
    try:
        await execute_insert(
            "INSERT INTO memo_exports (signal_id, format, file_path) VALUES (?, 'html', ?)",
            (signal_id, f"memo_{signal_id}_{datetime.now().strftime('%Y%m%d')}.html")
        )
    except Exception:
        pass

    return html


async def generate_memo_docx(signal_id: int) -> bytes | None:
    """시그널 데이터 기반 DOCX 투자 메모 생성"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    rows = await execute_query("SELECT * FROM signals WHERE id = ?", (signal_id,))
    if not rows:
        return None

    signal = dict(rows[0])
    scenarios = json.loads(signal.get("scenarios_json") or "{}")
    dart = json.loads(signal.get("dart_fundamentals_json") or "{}")
    expert_stances = json.loads(signal.get("expert_stances_json") or "{}")
    metadata = json.loads(signal.get("metadata_json") or "{}")
    direction_kr = {"buy": "매수", "sell": "매도", "hold": "보유"}.get(signal.get("direction", ""), signal.get("direction", ""))
    rr_score = signal.get("rr_score", 0) or 0

    doc = Document()

    # Title
    title = doc.add_heading(f'{signal.get("stock_name", "")} ({signal.get("stock_code", "")})', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f'생성일: {datetime.now().strftime("%Y-%m-%d %H:%M")} | 시그널 #{signal_id} | R/R Score: {rr_score:.1f}').font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Investment Opinion
    doc.add_heading('투자 의견', level=1)
    p = doc.add_paragraph()
    run = p.add_run(direction_kr)
    run.bold = True
    run.font.size = Pt(14)
    if signal.get("direction") == "buy":
        run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
    elif signal.get("direction") == "sell":
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    if signal.get("variant_view"):
        doc.add_paragraph(f'시장 오해: {signal["variant_view"]}')

    # Scenarios
    if scenarios:
        doc.add_heading('시나리오 분석', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '시나리오'
        hdr[1].text = '목표가'
        hdr[2].text = '상승률'
        hdr[3].text = '확률'
        for key in ["bull", "base", "bear"]:
            s = scenarios.get(key, {})
            if s:
                row = table.add_row().cells
                row[0].text = s.get("label", key)
                row[1].text = f'{s.get("price_target", 0):,.0f}원'
                row[2].text = f'{s.get("upside_pct", 0):+.1f}%'
                row[3].text = f'{s.get("probability", 0) * 100:.0f}%'

    # DART Fundamentals
    if dart:
        doc.add_heading('펀더멘탈 (DART)', level=1)
        labels = {
            "dart_per": "PER", "dart_pbr": "PBR", "dart_operating_margin": "영업이익률 (%)",
            "dart_debt_ratio": "부채비율 (%)", "dart_eps_yoy_pct": "EPS 성장률 (%)",
            "dart_dividend_yield": "배당수익률 (%)",
        }
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = '지표'
        table.rows[0].cells[1].text = '값'
        for key, label in labels.items():
            val = dart.get(key)
            if val is not None:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = f"{val:.1f}" if isinstance(val, (int, float)) else str(val)

    # Expert Panel
    if expert_stances:
        doc.add_heading('전문가 패널', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = '전문가'
        table.rows[0].cells[1].text = '의견'
        for name, stance in expert_stances.items():
            row = table.add_row().cells
            row[0].text = name
            row[1].text = stance

    # DCF
    dcf = metadata.get("dcf_valuation")
    if dcf and dcf.get("fair_value"):
        doc.add_heading('DCF 적정가', level=1)
        upside_text = f' ({dcf["upside_pct"]:+.1f}%)' if dcf.get("upside_pct") is not None else ''
        doc.add_paragraph(f'적정가: {dcf["fair_value"]:,.0f}원{upside_text}')

    # Investor trend
    investor = metadata.get("investor_trend")
    if investor:
        doc.add_heading(f'수급 동향 ({investor.get("days", 20)}일)', level=1)
        doc.add_paragraph(f'외국인 순매수: {investor.get("foreign_net_buy", 0):+,}주')
        doc.add_paragraph(f'기관 순매수: {investor.get("institution_net_buy", 0):+,}주')

    # Disclaimer
    doc.add_paragraph('')
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run('본 메모는 AI 분석 시스템에 의해 자동 생성되었습니다. 투자 판단의 참고 자료로만 활용하세요.')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # DB record
    try:
        await execute_insert(
            "INSERT INTO memo_exports (signal_id, format, file_path) VALUES (?, 'docx', ?)",
            (signal_id, f"memo_{signal_id}_{datetime.now().strftime('%Y%m%d')}.docx")
        )
    except Exception:
        pass

    return buffer.getvalue()
